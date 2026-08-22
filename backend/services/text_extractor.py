import re
from pathlib import Path
from pypdf import PdfReader
from backend.utils.logging_config import logger

def clean_extracted_text(text: str) -> str:
    """Clean whitespace while preserving paragraph line breaks."""
    if not text:
        return ""
    # Normalize Windows CRLF to LF
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    # Replace horizontal whitespace sequences (tabs, spaces) with a single space
    lines = [re.sub(r'[ \t]+', ' ', line).strip() for line in text.split("\n")]
    # Remove consecutive empty lines (allow max 1 blank line)
    cleaned_lines = []
    prev_empty = False
    for line in lines:
        if not line:
            if not prev_empty:
                cleaned_lines.append("")
                prev_empty = True
        else:
            cleaned_lines.append(line)
            prev_empty = False
    return "\n".join(cleaned_lines).strip()

def extract_text_from_txt(file_path: Path) -> str:
    """Extract text from TXT file with multiple encoding fallback."""
    encodings = ["utf-8", "latin-1", "cp1252", "iso-8859-1"]
    for enc in encodings:
        try:
            with open(file_path, "r", encoding=enc) as f:
                content = f.read()
                return clean_extracted_text(content)
        except UnicodeDecodeError:
            continue
    raise ValueError("Failed to decode text file with standard encodings (utf-8, latin-1).")

def extract_text_from_pdf(file_path: Path) -> str:
    """Extract machine-readable text from PDF file using pypdf."""
    try:
        reader = PdfReader(str(file_path))
        extracted_pages = []
        for i, page in enumerate(reader.pages):
            page_text = page.extract_text() or ""
            if page_text.strip():
                extracted_pages.append(page_text)
        
        full_text = "\n\n".join(extracted_pages)
        cleaned = clean_extracted_text(full_text)

        if not cleaned or len(cleaned) < 20:
            raise ValueError(
                "This PDF does not contain machine-readable text. Please upload a text-based PDF."
            )

        return cleaned
    except ValueError:
        raise
    except Exception as e:
        logger.error(f"Error reading PDF file {file_path}: {str(e)}")
        raise ValueError(f"Could not read PDF file: {str(e)}")

def extract_text_from_docx(file_path: Path) -> str:
    """Extract text from Microsoft Word .docx file using python-docx."""
    try:
        import docx
        doc = docx.Document(str(file_path))
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())
        
        # Also extract table text
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text.strip())

        text_content = "\n".join(full_text)
        cleaned = clean_extracted_text(text_content)

        if not cleaned or len(cleaned) < 10:
            raise ValueError("This Word document contains no extractable text.")

        return cleaned
    except Exception as e:
        logger.error(f"Error reading DOCX file {file_path}: {str(e)}")
        raise ValueError(f"Could not read Word document (.docx): {str(e)}")

def extract_text_from_file(file_path: Path) -> str:
    """Unified file text extractor route based on file extension."""
    ext = file_path.suffix.lower()
    if ext == ".txt":
        return extract_text_from_txt(file_path)
    elif ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext == ".docx":
        return extract_text_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file format: {ext}")
